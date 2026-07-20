"""
Solarix - DOCX Template Engine (Sprint 3)
- Extract {{ placeholder }} variables from any .docx
- Auto-suggest a mapping from placeholder → canonical system variable
- Render filled .docx by string-replacement (works with placeholders that contain spaces,
  which Jinja/docxtpl cannot handle natively).
"""
import re
import io
import unicodedata
import logging
from copy import deepcopy
from datetime import datetime, timezone
from typing import Dict, List, Tuple, Any, Optional

from docx import Document
import docx.document

logger = logging.getLogger(__name__)

PLACEHOLDER_RE = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")

# ----------------------------- System variable schema -----------------------------

# Each variable: key (snake_case canonical), label, source (client|company|computed|today|manual)
SYSTEM_VARIABLES: List[Dict[str, str]] = [
    # Client
    {"key": "client_full_name", "label": "Client Full Name", "source": "client", "path": "full_name"},
    {"key": "client_first_name", "label": "Client First Name", "source": "computed", "path": "full_name.split[0]"},
    {"key": "father_husband_name", "label": "Father/Husband Name", "source": "client", "path": "father_husband_name"},
    {"key": "mobile_no", "label": "Mobile Number", "source": "client", "path": "mobile"},
    {"key": "alt_mobile", "label": "Alternate Mobile", "source": "client", "path": "alt_mobile"},
    {"key": "email_id", "label": "Email", "source": "client", "path": "email"},
    {"key": "consumer_no", "label": "Consumer Number", "source": "client", "path": "consumer_number"},
    {"key": "consumer_type", "label": "Consumer Type (LT/HT)", "source": "computed", "path": "phase_type→LT/HT"},
    {"key": "address", "label": "Installation Address", "source": "client", "path": "address"},
    {"key": "village", "label": "Village", "source": "client", "path": "village"},
    {"key": "taluka", "label": "Taluka", "source": "client", "path": "taluka"},
    {"key": "district", "label": "District", "source": "client", "path": "district"},
    {"key": "city", "label": "City", "source": "client", "path": "city"},
    {"key": "state", "label": "State", "source": "client", "path": "state"},
    {"key": "pincode", "label": "Pincode", "source": "client", "path": "pincode"},
    {"key": "aadhaar_no", "label": "Aadhaar Number", "source": "client", "path": "aadhaar"},
    {"key": "pan", "label": "PAN", "source": "client", "path": "pan"},
    {"key": "load", "label": "Load", "source": "client", "path": "load"},
    {"key": "sanction_load", "label": "Sanction Load", "source": "client", "path": "sanction_load"},
    # Solar system
    {"key": "panel_in_kw", "label": "System Size (kW)", "source": "client", "path": "system_kw"},
    {"key": "panel_brand", "label": "Panel Brand", "source": "client", "path": "panel_make"},
    {"key": "panel_in_wp", "label": "Panel Wattage (Wp)", "source": "client", "path": "panel_wattage"},
    {"key": "panel_type", "label": "Panel Type / Model", "source": "computed", "path": "panel_make + panel_wattage"},
    {"key": "panel_brand_and_panel_type", "label": "Panel Brand + Type", "source": "computed", "path": "panel_make + panel_wattage"},
    {"key": "solar_panels_in_nos", "label": "Number of Panels", "source": "client", "path": "num_panels"},
    {"key": "panel_serial_numbers", "label": "Panel Serial Numbers", "source": "client", "path": "panel_serial_numbers"},
    # Inverter
    {"key": "inverter_brand", "label": "Inverter Brand", "source": "client", "path": "inverter_make"},
    {"key": "inverter_model", "label": "Inverter Model", "source": "client", "path": "inverter_model"},
    {"key": "inverter_in_kw", "label": "Inverter Capacity (kW)", "source": "client", "path": "inverter_capacity"},
    {"key": "inverter_serial_no", "label": "Inverter Serial Number", "source": "client", "path": "inverter_serial"},
    {"key": "inverter_nos", "label": "Number of Inverters", "source": "computed", "path": "fixed=1 unless override"},
    {"key": "inverter_kw_times_inverter_nos", "label": "Inverter kW × Nos", "source": "computed", "path": "inverter_in_kw + inverter_nos"},
    {"key": "solar_wp_times_panel_nos", "label": "Solar Wp × Panel Nos", "source": "computed", "path": "panel_in_wp × solar_panels_in_nos"},
    # Project / commissioning
    {"key": "net_meter_number", "label": "Net Meter Number", "source": "client", "path": "net_meter_number"},
    {"key": "meter_number", "label": "Meter Number", "source": "client", "path": "meter_number"},
    {"key": "vendor_name", "label": "Vendor / Company Name", "source": "company", "path": "company_name"},
    {"key": "installer", "label": "Installer", "source": "client", "path": "installer"},
    {"key": "bu_no", "label": "BU Number", "source": "manual", "path": "manual"},
    {"key": "survey_date", "label": "Survey Date", "source": "client", "path": "survey_date"},
    {"key": "installation_date", "label": "Installation Date", "source": "computed", "path": "today"},
    {"key": "latitude", "label": "Latitude", "source": "client", "path": "latitude"},
    {"key": "longitude", "label": "Longitude", "source": "client", "path": "longitude"},
    {"key": "roof_type", "label": "Roof Type", "source": "client", "path": "roof_type"},
    {"key": "project_id", "label": "Project ID", "source": "client", "path": "project_id"},
    {"key": "application_number", "label": "Application Number", "source": "client", "path": "application_number"},
    # Company / Vendor
    {"key": "vendor_address", "label": "Vendor Address", "source": "company", "path": "address"},
    {"key": "vendor_office_address", "label": "Vendor Office Address", "source": "company", "path": "address"},
    {"key": "vendor_mobile", "label": "Vendor Mobile", "source": "company", "path": "mobile"},
    {"key": "vendor_email", "label": "Vendor Email", "source": "company", "path": "email"},
    {"key": "vendor_gst", "label": "Vendor GST", "source": "company", "path": "gst_number"},
    # Today
    {"key": "date", "label": "Today's Date (DD-MM-YYYY)", "source": "today", "path": "today"},
    {"key": "day", "label": "Day of Month", "source": "today", "path": "today.day"},
    {"key": "month", "label": "Month Name", "source": "today", "path": "today.month"},
    {"key": "year", "label": "Year", "source": "today", "path": "today.year"},
]

SYSTEM_VAR_KEYS = {v["key"] for v in SYSTEM_VARIABLES}


# ----------------------------- Placeholder extraction -----------------------------

def _iter_paragraphs(doc: docx.document.Document):
    yield from doc.paragraphs
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def extract_placeholders(file_bytes: bytes) -> List[str]:
    """Return list of unique placeholder names (without {{ }}), preserving discovery order."""
    doc = Document(io.BytesIO(file_bytes))
    seen: List[str] = []
    seen_set = set()
    # Headers / footers
    sections = list(doc.sections)
    sources: List[Any] = [doc]
    for s in sections:
        sources.extend([s.header, s.footer])
    # Collect all text from paragraphs and table cells
    def walk(obj):
        if hasattr(obj, "paragraphs"):
            for p in obj.paragraphs:
                _scan_text(p.text, seen, seen_set)
        if hasattr(obj, "tables"):
            for tbl in obj.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        walk(cell)
    for src in sources:
        walk(src)
    return seen


def _scan_text(text: str, seen: List[str], seen_set: set):
    for m in PLACEHOLDER_RE.finditer(text or ""):
        name = m.group(1).strip()
        # collapse multiple whitespace
        name = re.sub(r"\s+", " ", name)
        if name not in seen_set:
            seen.append(name)
            seen_set.add(name)


# ----------------------------- Smart auto-mapping -----------------------------

def _normalize(s: str) -> str:
    s = s.lower().strip()
    s = unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode()
    s = re.sub(r"[^a-z0-9 ]+", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


# canonical_var_key -> list of search tokens (lowercase, normalised)
MAPPING_HINTS: Dict[str, List[str]] = {
    "client_full_name": ["client full name", "client name", "name of consumer", "customer name", "consumer name", "applicant name", "client"],
    "father_husband_name": ["father husband name", "father name", "husband name", "father husband name", "father/husband name"],
    "mobile_no": ["mobile no", "mobile number", "phone", "contact number", "mobile"],
    "alt_mobile": ["alt mobile", "alternate mobile", "alternate mobile number"],
    "email_id": ["email id", "email", "e mail", "mail id"],
    "consumer_no": ["consumer no", "consumer number", "consumer id"],
    "address": ["address", "addre ss", "installation address", "client address", "clionet address", "site address", "premises"],
    "village": ["village", "town"],
    "taluka": ["taluka", "tehsil", "block"],
    "district": ["district"],
    "city": ["city"],
    "state": ["state"],
    "pincode": ["pincode", "pin code"],
    "aadhaar_no": ["aadhaar no", "aadhar no", "aadhaar photo", "aaddhar photo", "aadhar photo", "uid"],
    "pan": ["pan", "pan number", "pan card"],
    "load": ["load", "connected load"],
    "sanction_load": ["sanction load", "sanctioned load"],
    "consumer_type": ["consumer type", "category", "lt ht"],
    "panel_in_kw": ["panel in kw", "solar in kw", "solar kw", "system in kw", "kw", "rooftop capacity", "panel kw", "re installed capacity"],
    "panel_brand": ["panel brand", "panel make", "module brand", "brand name", "brand", "module make"],
    "panel_type": ["panel type", "panel model", "module type"],
    "panel_brand_and_panel_type": ["panel brand and panel type", "panel brand and type", "panel brand type", "module brand and type"],
    "panel_in_wp": ["solar in wp", "solar wp", "panel wp", "module capacity", "module wattage", "wp"],
    "solar_panels_in_nos": ["solar panels in nos", "solar in nos", "no of panels", "no of modules", "panel nos", "panels", "modules"],
    "panel_serial_numbers": ["panel serial numbers", "panel serials", "panel serial no"],
    "inverter_brand": ["inverter brand", "inverter make", "brand of inverter", "make of inverter"],
    "inverter_model": ["inverter model", "inverter type"],
    "inverter_in_kw": ["inverter in kw", "inverter kw", "inverter capacity"],
    "inverter_serial_no": ["inverter serial no", "inverter serial number", "inverter sl no", "inverter sn"],
    "inverter_nos": ["inverter nos", "no of inverter", "inverter count"],
    "inverter_kw_times_inverter_nos": ["inverter kw inverter nos", "inverter kw x nos", "inverter kw nos"],
    "solar_wp_times_panel_nos": ["solar wp panel nos", "solar wp panels nos", "solar wp panle nos", "wp x panels"],
    "net_meter_number": ["net meter number", "net meter no"],
    "meter_number": ["meter number", "meter no"],
    "vendor_name": ["vendor name", "company name", "epc name", "vendor"],
    "installer": ["installer", "installer name"],
    "bu_no": ["bu no", "bu number", "business unit"],
    "survey_date": ["survey date", "surveyed date"],
    "installation_date": ["installation date", "commissioning date", "install date"],
    "latitude": ["latitude", "lat"],
    "longitude": ["longitude", "lng", "lon"],
    "roof_type": ["roof type", "roofing"],
    "project_id": ["project id", "sol id", "project number"],
    "application_number": ["application number", "application no", "appl no"],
    "vendor_address": ["vendor address", "vendor office address", "office address", "company address", "vendor address office address"],
    "vendor_mobile": ["vendor mobile", "vendor phone", "company mobile"],
    "vendor_email": ["vendor email", "company email"],
    "vendor_gst": ["vendor gst", "gst number", "gstin"],
    "date": ["date"],
    "day": ["day"],
    "month": ["month"],
    "year": ["year"],
}


def suggest_mapping(placeholders: List[str]) -> Dict[str, str]:
    """Given list of raw placeholder strings, return {placeholder: system_var_key or ''}."""
    result: Dict[str, str] = {}
    for raw in placeholders:
        norm = _normalize(raw)
        best_key = ""
        best_score = 0
        for key, hints in MAPPING_HINTS.items():
            for hint in hints:
                hnorm = _normalize(hint)
                if not hnorm:
                    continue
                # exact match
                if norm == hnorm:
                    score = 1000 + len(hnorm)
                # substring match
                elif hnorm in norm or norm in hnorm:
                    score = 100 + len(hnorm) - abs(len(hnorm) - len(norm))
                else:
                    # token overlap
                    a = set(norm.split())
                    b = set(hnorm.split())
                    if not a or not b:
                        continue
                    overlap = a & b
                    if not overlap:
                        continue
                    score = (len(overlap) * 20) - abs(len(a) - len(b))
                if score > best_score:
                    best_score = score
                    best_key = key
        result[raw] = best_key if best_score >= 40 else ""
    return result


# ----------------------------- Value resolution -----------------------------

def _safe_str(v: Any) -> str:
    if v is None:
        return ""
    if isinstance(v, float):
        # avoid 5.0 → "5.0" when integer
        return str(int(v)) if v.is_integer() else f"{v:.2f}".rstrip("0").rstrip(".")
    return str(v)


def resolve_value(key: str, client: dict, company: dict, overrides: Dict[str, Any]) -> str:
    """Resolve canonical system variable key to a string value, applying overrides if present."""
    if key in overrides and overrides[key] not in (None, ""):
        return _safe_str(overrides[key])

    client = client or {}
    company = company or {}
    now = datetime.now(timezone.utc)

    if key == "client_full_name":
        return _safe_str(client.get("full_name") or client.get("client_name"))
    if key == "client_first_name":
        name = client.get("full_name") or client.get("client_name") or ""
        return _safe_str(name.split()[0] if name else "")
    if key == "father_husband_name":
        return _safe_str(client.get("father_husband_name") or client.get("father_name") or client.get("husband_name") or client.get("father") or client.get("husband"))
    if key == "mobile_no":
        return _safe_str(client.get("mobile") or client.get("mobile_no"))
    if key == "alt_mobile":
        return _safe_str(client.get("alt_mobile") or client.get("alternate_mobile"))
    if key == "email_id":
        return _safe_str(client.get("email") or client.get("email_id"))
    if key == "consumer_no":
        return _safe_str(client.get("consumer_number") or client.get("consumer_no"))
    if key == "address":
        if client.get("address"):
            return _safe_str(client.get("address"))
        parts = [client.get("address"), client.get("city"), client.get("state"), client.get("pincode")]
        return ", ".join([p for p in parts if p])
    if key == "village":
        return _safe_str(client.get("village") or client.get("city") or client.get("town"))
    if key == "taluka":
        return _safe_str(client.get("taluka") or client.get("sub_district") or client.get("block"))
    if key == "district":
        return _safe_str(client.get("district"))
    if key == "city":
        return _safe_str(client.get("city") or client.get("village"))
    if key == "state":
        return _safe_str(client.get("state"))
    if key == "pincode":
        return _safe_str(client.get("pincode") or client.get("pin_code"))
    if key == "aadhaar_no":
        return _safe_str(client.get("aadhaar") or client.get("aadhaar_no") or client.get("aadhar"))
    if key == "pan":
        return _safe_str(client.get("pan") or client.get("pan_no") or client.get("pan_number"))
    if key == "load":
        return _safe_str(client.get("load") or client.get("load_kw") or client.get("connected_load") or client.get("sanction_load") or client.get("sanctioned_load"))
    if key == "sanction_load":
        return _safe_str(client.get("sanction_load") or client.get("sanctioned_load") or client.get("load") or client.get("connected_load"))
    if key == "consumer_type":
        ph = (client.get("phase_type") or client.get("consumer_type") or "").lower()
        if "three" in ph or "3" in ph or "ht" in ph:
            return "HT"
        return "LT"
    if key == "panel_in_kw":
        return _safe_str(client.get("system_kw") or client.get("solar_capacity") or client.get("capacity"))
    if key == "panel_brand":
        return _safe_str(client.get("panel_make") or client.get("panel_brand"))
    if key == "panel_in_wp":
        return _safe_str(client.get("panel_wattage") or client.get("panel_watt") or client.get("panel_wp"))
    if key == "panel_type":
        wp = client.get("panel_wattage") or client.get("panel_watt") or client.get("panel_wp")
        return f"{wp} Wp Mono PERC" if wp else ""
    if key == "panel_brand_and_panel_type":
        brand = client.get("panel_make") or client.get("panel_brand") or ""
        wp = client.get("panel_wattage") or client.get("panel_watt") or client.get("panel_wp")
        if brand and wp:
            return f"{brand} {wp}Wp Mono PERC"
        return brand or (f"{wp}Wp" if wp else "")
    if key == "solar_panels_in_nos":
        return _safe_str(client.get("num_panels") or client.get("number_of_panels"))
    if key == "panel_serial_numbers":
        return _safe_str(client.get("panel_serial_numbers") or client.get("panel_serials") or client.get("panel_serial_no") or client.get("panel_serial"))
    if key == "inverter_brand":
        return _safe_str(client.get("inverter_make") or client.get("inverter_brand"))
    if key == "inverter_model":
        return _safe_str(client.get("inverter_model") or client.get("inverter_type"))
    if key == "inverter_in_kw":
        return _safe_str(client.get("inverter_capacity") or client.get("inverter_kw"))
    if key == "inverter_serial_no":
        return _safe_str(client.get("inverter_serial") or client.get("inverter_serial_no") or client.get("inverter_serial_numbers"))
    if key == "inverter_nos":
        return _safe_str(overrides.get("inverter_nos") or client.get("inverter_nos") or 1)
    if key == "inverter_kw_times_inverter_nos":
        ikw = client.get("inverter_capacity") or client.get("inverter_kw") or ""
        n = overrides.get("inverter_nos") or client.get("inverter_nos") or 1
        return f"{ikw} × {n}" if ikw else ""
    if key == "solar_wp_times_panel_nos":
        wp = client.get("panel_wattage") or client.get("panel_watt") or client.get("panel_wp") or 0
        n = client.get("num_panels") or client.get("number_of_panels") or 0
        if wp and n:
            try:
                total_kw = (float(wp) * float(n)) / 1000.0
                return f"{int(wp)}Wp × {int(n)} = {total_kw:g} kW"
            except Exception:
                return f"{wp}Wp × {n}"
        return ""
    if key == "net_meter_number":
        return _safe_str(client.get("net_meter_number") or client.get("net_meter_no") or client.get("net_meter"))
    if key == "meter_number":
        return _safe_str(client.get("meter_number") or client.get("meter_no") or client.get("meter"))
    if key == "vendor_name":
        return _safe_str(client.get("vendor") or client.get("vendor_name") or company.get("company_name"))
    if key == "installer":
        return _safe_str(client.get("installer") or client.get("installer_name"))
    if key == "bu_no":
        return _safe_str(overrides.get("bu_no") or client.get("bu_no"))
    if key == "survey_date":
        return _safe_str(client.get("survey_date") or client.get("surveyed_date"))
    if key == "installation_date":
        return _safe_str(overrides.get("installation_date") or client.get("install_date") or client.get("installation_date") or now.strftime("%d-%m-%Y"))
    if key == "latitude":
        return _safe_str(client.get("latitude") or client.get("lat"))
    if key == "longitude":
        return _safe_str(client.get("longitude") or client.get("lng") or client.get("lon"))
    if key == "roof_type":
        return _safe_str(client.get("roof_type") or client.get("roof"))
    if key == "project_id":
        return _safe_str(client.get("sol_id") or client.get("project_id"))
    if key == "application_number":
        return _safe_str(client.get("application_number") or client.get("application_no") or client.get("appl_no"))
    if key == "vendor_address":
        parts = [company.get("address"), company.get("city"), company.get("state"), company.get("pincode")]
        return ", ".join([p for p in parts if p])
    if key == "vendor_office_address":
        parts = [company.get("address"), company.get("city"), company.get("state"), company.get("pincode")]
        return ", ".join([p for p in parts if p])
    if key == "vendor_mobile":
        return _safe_str(company.get("mobile"))
    if key == "vendor_email":
        return _safe_str(company.get("email"))
    if key == "vendor_gst":
        return _safe_str(company.get("gst_number"))
    if key == "date":
        return now.strftime("%d-%m-%Y")
    if key == "day":
        return str(now.day)
    if key == "month":
        return now.strftime("%B")
    if key == "year":
        return str(now.year)
    return ""


# ----------------------------- Run-aware DOCX substitution -----------------------------

def _replace_in_paragraph(paragraph, replacements: Dict[str, str]):
    """
    Replace each `{{ name }}` -> value while preserving the formatting of the first run that
    starts the placeholder. Word frequently splits {{name}} across multiple runs; we therefore
    join all runs, replace, then write the result back into the first run and clear the rest.
    """
    runs = paragraph.runs
    if not runs:
        return
    full = "".join(r.text for r in runs)
    if "{{" not in full:
        return
    new_text = full
    for key, val in replacements.items():
        # tokenize the canonical key on whitespace and rebuild a regex that tolerates any
        # amount of whitespace between the original tokens (Word often has double/non-breaking spaces)
        tokens = key.split()
        if not tokens:
            continue
        body = r"\s+".join(re.escape(t) for t in tokens)
        pattern = re.compile(r"\{\{\s*" + body + r"\s*\}\}")
        new_text = pattern.sub(lambda m: val, new_text)
    if new_text == full:
        return
    # write back: keep first run formatting, clear others
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def _walk_and_replace(container, replacements: Dict[str, str]):
    if hasattr(container, "paragraphs"):
        for p in container.paragraphs:
            _replace_in_paragraph(p, replacements)
    if hasattr(container, "tables"):
        for tbl in container.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    _walk_and_replace(cell, replacements)


# ----------------------------- Schema-less Fallback Lookup -----------------------------

ALIASES: Dict[str, List[str]] = {
    "client_full_name": ["full_name", "client_full_name", "client_name", "name"],
    "client_name": ["full_name", "client_full_name", "client_name", "name"],
    "father_husband_name": ["father_husband_name", "father_name", "husband_name", "father", "husband"],
    "address": ["address", "installation_address"],
    "village": ["village", "city", "town"],
    "taluka": ["taluka", "sub_district", "block"],
    "district": ["district"],
    "state": ["state"],
    "pincode": ["pincode", "pin_code", "pin"],
    "pin_code": ["pincode", "pin_code", "pin"],
    "consumer_number": ["consumer_number", "consumer_no", "consumer_id"],
    "consumer_no": ["consumer_number", "consumer_no", "consumer_id"],
    "consumer_type": ["consumer_type", "phase_type", "phase"],
    "mobile": ["mobile", "mobile_no", "phone", "phone_number"],
    "mobile_no": ["mobile", "mobile_no", "phone", "phone_number"],
    "alternate_mobile": ["alt_mobile", "alternate_mobile"],
    "alt_mobile": ["alt_mobile", "alternate_mobile"],
    "email": ["email", "email_id", "email_address"],
    "email_id": ["email", "email_id", "email_address"],
    "aadhaar": ["aadhaar", "aadhar", "aadhaar_no", "aadhar_no", "aadhaar_number", "aadhar_number"],
    "aadhaar_no": ["aadhaar", "aadhar", "aadhaar_no", "aadhar_no", "aadhaar_number", "aadhar_number"],
    "pan": ["pan", "pan_no", "pan_number"],
    "load": ["load", "load_kw", "connected_load", "sanction_load", "sanctioned_load"],
    "sanction_load": ["sanction_load", "sanctioned_load", "load", "connected_load"],
    "solar_capacity": ["system_kw", "solar_capacity", "capacity", "solar_capacity_kw"],
    "solar_capacity_kw": ["system_kw", "solar_capacity", "capacity", "solar_capacity_kw"],
    "panel_in_kw": ["system_kw", "solar_capacity", "capacity", "solar_capacity_kw"],
    "panel_brand": ["panel_make", "panel_brand", "panel_manufacturer"],
    "panel_watt": ["panel_wattage", "panel_watt", "panel_wp", "module_wattage"],
    "panel_in_wp": ["panel_wattage", "panel_watt", "panel_wp", "module_wattage"],
    "number_of_panels": ["num_panels", "number_of_panels", "panel_count"],
    "solar_panels_in_nos": ["num_panels", "number_of_panels", "panel_count"],
    "panel_serial_numbers": ["panel_serial_numbers", "panel_serials", "panel_serial_no", "panel_serial"],
    "inverter_brand": ["inverter_make", "inverter_brand", "inverter_manufacturer"],
    "inverter_model": ["inverter_model", "inverter_type"],
    "inverter_capacity": ["inverter_capacity", "inverter_kw", "inverter_size"],
    "inverter_in_kw": ["inverter_capacity", "inverter_kw", "inverter_size"],
    "inverter_serial_numbers": ["inverter_serial", "inverter_serial_no", "inverter_serials", "inverter_serial_numbers"],
    "inverter_serial_no": ["inverter_serial", "inverter_serial_no", "inverter_serials", "inverter_serial_numbers"],
    "net_meter_number": ["net_meter_number", "net_meter_no", "net_meter"],
    "meter_number": ["meter_number", "meter_no", "meter"],
    "vendor": ["vendor", "vendor_name", "company_name"],
    "vendor_name": ["vendor", "vendor_name", "company_name"],
    "installer": ["installer", "installer_name"],
    "survey_date": ["survey_date", "surveyed_date"],
    "installation_date": ["install_date", "installation_date", "commissioning_date"],
    "latitude": ["latitude", "lat"],
    "longitude": ["longitude", "lng", "lon"],
    "roof_type": ["roof_type", "roof"],
    "project_id": ["sol_id", "project_id"],
    "application_number": ["application_number", "application_no", "appl_no"],
}

def normalize_key(k: str) -> str:
    k = k.lower().strip()
    k = re.sub(r"[^a-z0-9]", "_", k)
    k = re.sub(r"_+", "_", k)
    return k.strip("_")

def resolve_client_field(placeholder: str, client: dict) -> Optional[Any]:
    if not client:
        return None
    ph_clean = placeholder.strip()
    # 1. Direct case-insensitive key check
    for k, v in client.items():
        if k.lower() == ph_clean.lower():
            return v
    # 2. Normalize and check direct match
    norm_ph = normalize_key(placeholder)
    for k, v in client.items():
        if normalize_key(k) == norm_ph:
            return v
    # 3. Check aliases
    aliases = ALIASES.get(norm_ph, [])
    for alias in aliases:
        norm_alias = normalize_key(alias)
        for k, v in client.items():
            if normalize_key(k) == norm_alias:
                return v
    return None

def render_docx(file_bytes: bytes,
                placeholders: List[str],
                mapping: Dict[str, str],
                client: dict,
                company: dict,
                overrides: Optional[Dict[str, Any]] = None) -> bytes:
    """
    Open .docx and replace every placeholder. `mapping[placeholder]` = system_var_key (or "").
    For unmapped placeholders, use overrides["__raw__"][placeholder] if provided, else blank.
    """
    overrides = overrides or {}
    raw_overrides = overrides.get("__raw__", {})

    replacements: Dict[str, str] = {}
    for ph in placeholders:
        # raw_overrides always win (lets the user edit any field in the generation dialog,
        # even when it's auto-mapped to a canonical variable)
        if ph in raw_overrides and raw_overrides[ph] not in (None, ""):
            replacements[ph] = _safe_str(raw_overrides[ph])
            continue
        var_key = mapping.get(ph, "") if mapping else ""
        val = ""
        if var_key and var_key in SYSTEM_VAR_KEYS:
            val = resolve_value(var_key, client, company, overrides)
        elif ph in raw_overrides:  # explicit empty override
            val = _safe_str(raw_overrides[ph])
        
        # Fallback to schema-less lookup if still empty
        if not val:
            fallback_val = resolve_client_field(ph, client)
            if fallback_val is not None:
                val = _safe_str(fallback_val)
        replacements[ph] = val

    doc = Document(io.BytesIO(file_bytes))
    _walk_and_replace(doc, replacements)
    # Headers / footers
    for s in doc.sections:
        _walk_and_replace(s.header, replacements)
        _walk_and_replace(s.footer, replacements)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ----------------------------- Preview values (for UI) -----------------------------

def compute_preview(placeholders: List[str],
                    mapping: Dict[str, str],
                    client: dict,
                    company: dict,
                    overrides: Optional[Dict[str, Any]] = None) -> List[Dict[str, str]]:
    """Return per-placeholder preview rows for the generation dialog."""
    overrides = overrides or {}
    raw_overrides = overrides.get("__raw__", {})
    out = []
    for ph in placeholders:
        var_key = (mapping or {}).get(ph, "")
        # raw_overrides always win — gives users a single edit surface in the UI
        if ph in raw_overrides and raw_overrides[ph] not in (None, ""):
            val = _safe_str(raw_overrides[ph])
            source = (next((v["label"] for v in SYSTEM_VARIABLES if v["key"] == var_key), var_key) if var_key
                      else "Manual override") + " (manual)"
            out.append({
                "placeholder": ph, "variable": var_key,
                "label": source, "value": val, "missing": False,
            })
            continue
        
        val = ""
        source = ""
        if var_key and var_key in SYSTEM_VAR_KEYS:
            val = resolve_value(var_key, client, company, overrides)
            source = next((v["label"] for v in SYSTEM_VARIABLES if v["key"] == var_key), var_key)
        elif ph in raw_overrides:
            val = _safe_str(raw_overrides[ph])
            source = "Manual override"
        
        # Fallback to schema-less lookup if still empty
        if not val:
            fallback_val = resolve_client_field(ph, client)
            if fallback_val is not None:
                val = _safe_str(fallback_val)
                source = "Auto-resolved field"
        
        if not source:
            source = "— unmapped —"
            
        out.append({
            "placeholder": ph, "variable": var_key,
            "label": source, "value": val, "missing": not val,
        })
    return out
